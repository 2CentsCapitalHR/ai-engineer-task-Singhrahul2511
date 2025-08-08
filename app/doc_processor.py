# app/doc_processor.py

import docx
import io
import re
from typing import Tuple, List, Dict
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

from config import DOCUMENT_KEYWORDS
from app.utils import logger

def parse_docx(file_upload) -> Tuple[str, str]:
    """
    Parses an uploaded .docx file to extract its text content.
    """
    try:
        file_stream = io.BytesIO(file_upload.getvalue())
        doc = docx.Document(file_stream)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        logger.info(f"Successfully parsed document: {file_upload.name}")
        return full_text, None
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_upload.name}: {e}", exc_info=True)
        return None, f"Failed to parse {file_upload.name}. It might be corrupted."

def identify_document_type(doc_text: str) -> str:
    """
    Identifies the document type based on keywords.
    """
    doc_text_lower = doc_text.lower()
    for doc_type, keywords in DOCUMENT_KEYWORDS.items():
        if any(keyword in doc_text_lower for keyword in keywords):
            logger.info(f"Document classified as: {doc_type}")
            return doc_type
    logger.warning("Could not classify document type.")
    return "Unknown"

def add_comments_to_docx(file_stream: io.BytesIO, issues: List[Dict[str, str]]) -> io.BytesIO:
    """
    Adds multiple formatted comment paragraphs to a .docx file based on a list of issues.
    This version uses more robust matching to place comments accurately.
    """
    doc = docx.Document(file_stream)
    
    commented_sections = set()

    # --- Handle General/Overall Comments ---
    general_comments = [
        issue for issue in issues 
        if any(keyword in issue.get("section", "").lower() for keyword in ["document", "overall", "general"])
    ]
    
    if general_comments:
        first_paragraph = doc.paragraphs[0]
        for issue in reversed(general_comments):
            section_key = issue.get("section", "").lower().strip()
            if section_key in commented_sections: continue
            
            comment_text = f"Issue: {issue.get('issue')}\nSuggestion: {issue.get('suggestion')}"
            comment_p = first_paragraph.insert_paragraph_before(f'[Corporate Agent General Comment]: {comment_text}')
            for run in comment_p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xB2, 0x22, 0x22)
                run.font.size = Pt(10)
            logger.info(f"Added general comment for section: '{issue.get('section')}'")
            commented_sections.add(section_key)

    # --- Process Specific Section Comments ---
    for issue in issues:
        original_section_text = issue.get("section", "")
        
        # --- THE FIX IS HERE ---
        # Clean the section text from the AI to improve matching.
        # This removes things like "Clause X:", "Section Y.", commas, etc.
        cleaned_section_text = re.sub(r'^(clause|section)\s*\d*[,.]?\s*', '', original_section_text, flags=re.IGNORECASE).strip().lower()
        # Further clean common trailing words
        cleaned_section_text = re.sub(r'\s*(clause|section)$', '', cleaned_section_text, flags=re.IGNORECASE).strip()

        if not cleaned_section_text or cleaned_section_text in commented_sections:
            continue

        comment_text = f"Issue: {issue.get('issue')}\nSuggestion: {issue.get('suggestion')}"

        target_para = None
        for p in doc.paragraphs:
            # Check if the cleaned section text is a substring of the paragraph's text
            if cleaned_section_text in p.text.lower().strip():
                target_para = p
                break
        
        if target_para:
            for r in target_para.runs:
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW

            comment_p = target_para.insert_paragraph_before(f'[Corporate Agent Comment]: {comment_text}')
            
            for run in comment_p.runs:
                 run.font.italic = True
                 run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
                 run.font.size = Pt(9)
            
            logger.info(f"Added comment for section: '{original_section_text}'")
            commented_sections.add(cleaned_section_text)
        else:
            logger.warning(f"Could not find a matching paragraph for section: '{original_section_text}'")

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream
